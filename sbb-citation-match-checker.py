import re
from pathlib import Path

from docx import Document


def extract_text_from_docx(path: Path) -> str:
    """
    Extract all text from a .docx file, including paragraphs and table cells.
    """
    doc = Document(path)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)


def extract_citations(text: str):
    """
    Extract in-text citation candidates in name–year style.

    Rule: any parentheses content that contains at least one 4-digit year
    (optionally with a, b, c suffix) is treated as a citation candidate.

    Examples:
      (Duer et al., 1992)
      (Duer et al., 1992; Luo, 2022)
      (Allan, 1999, Allan and Jones, 1999, Allan 2000a, Allan 2000b)
    """
    pattern = re.compile(
        r"\("                            # opening parenthesis
        r"[^()]*?"                       # any content inside (non-greedy)
        r"\d{4}[a-z]?[^()]*?"            # at least one year (e.g. 1992, 2000a)
        r"(?:[;,][^()]*\d{4}[a-z]?[^()]*)*"  # optionally more items ; or , + year
        r"\)"                            # closing parenthesis
    )

    citations_raw = [m.group(0) for m in pattern.finditer(text)]
    citations_clean = [c[1:-1].strip() for c in citations_raw]  # strip parentheses

    return citations_clean


def normalize_author(author: str) -> str:
    """
    Normalize an author string to a 'surname' key for matching.

    Rules:
    - Remove 'et al.'
    - If format is 'Surname, Initials', keep 'Surname'
    - If format is 'X. Li', treat the last token as surname (Li)
    - Otherwise, use the first token as surname
    - Strip non-letter characters except hyphen, apostrophe and spaces
    - Return lowercase surname
    """
    author = author.strip()

    # Remove 'et al.' (permissive so it catches 'X. Li et al.')
    author = re.sub(r"et al\.", "", author).strip()

    # Handle 'Surname, Initials'
    if "," in author:
        author = author.split(",", 1)[0].strip()

    tokens = author.split()
    if not tokens:
        return ""

    # Handle 'X. Li' type: leading initial + surname
    if len(tokens) >= 2 and re.match(r"^[A-Z]\.?$", tokens[0]):
        surname = tokens[-1]
    else:
        # Default: first token is surname
        surname = tokens[0]

    # Remove strange characters from surname (keep letters, hyphen, apostrophe, space)
    surname = re.sub(r"[^A-Za-z\-'\s]", "", surname)

    return surname.lower()


def build_citation_keys(citation_strings):
    """
    Parse all in-text citation strings and build keys of the form:
      '<normalized_first_author>|<year>'
    where year may include a/b/c suffix.

    Returns:
      dict: key -> set of original citation snippets where this key appeared
    """
    pattern = re.compile(r"([A-Z][^,]*?),\s*(\d{4}[a-z]?)")

    key_to_examples = {}

    for cit in citation_strings:
        for m in pattern.finditer(cit):
            raw_author_part = m.group(1).strip()
            year = m.group(2)

            author_norm = normalize_author(raw_author_part)
            if not author_norm:
                continue

            key = f"{author_norm}|{year}"
            key_to_examples.setdefault(key, set()).add(cit)

    return key_to_examples


def extract_reference_paragraphs(doc: Document):
    """
    Extract paragraphs that belong to the reference list.

    Assumptions:
    - There is a heading containing one of several keywords:
      'references', 'reference', 'bibliography', 'literature cited',
      'works cited', '参考文献', '參考文獻', '文献', '文獻'.
    - All non-empty paragraphs after that heading are treated as reference entries.
    """
    heading_keywords = [
        "references",
        "reference",
        "bibliography",
        "literature cited",
        "works cited",
        "参考文献",
        "參考文獻",
        "文献",
        "文獻",
    ]

    refs_started = False
    refs = []

    for p in doc.paragraphs:
        text = p.text.strip()
        text_lower = text.lower()

        if not refs_started:
            if any(kw in text_lower for kw in heading_keywords):
                refs_started = True
            continue

        if not text:
            # skip empty lines
            continue

        refs.append(text)

    return refs


def parse_reference_entries(ref_lines):
    """
    Parse reference entries and build keys of the form:
      '<normalized_first_author>|<year>'

    Assumptions:
    - The first author is before the first comma.
    - The first year in the line is the publication year (with optional a/b).
    """
    entries = []

    for i, text in enumerate(ref_lines):
        # Find year
        m = re.search(r"(\d{4}[a-z]?)", text)
        if not m:
            # No year detected; skip
            continue

        year = m.group(1)

        # First author: text before the first comma
        first_author_part = text.split(",", 1)[0].strip()
        author_norm = normalize_author(first_author_part)
        if not author_norm:
            continue

        key = f"{author_norm}|{year}"

        entries.append(
            {
                "index": i,
                "raw": text,
                "author_norm": author_norm,
                "year": year,
                "key": key,
            }
        )

    return entries


def main():
    # Path to your Word file
    docx_path = Path(r"C:\path\to\your\docx\file.docx")
    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    # Output paths
    citations_txt_path = docx_path.with_name(docx_path.stem + "_citations.txt")
    report_txt_path = docx_path.with_name(docx_path.stem + "_citation_report.txt")

    log_lines = []  # all text to be written to the report

    # ------------------------------------------------------------------
    # 1) Read full text and extract in-text citations
    # ------------------------------------------------------------------
    log_lines.append(f"DOCX file: {docx_path}")
    full_text = extract_text_from_docx(docx_path)

    log_lines.append("\n[Step 1] Extracting in-text citation candidates...")
    citations = extract_citations(full_text)
    log_lines.append(f"Number of citation parentheses found: {len(citations)}")

    # Save raw citation parentheses (for manual inspection if desired)
    with citations_txt_path.open("w", encoding="utf-8") as f:
        for c in citations:
            f.write(c + "\n")
    log_lines.append(f"Citation parentheses saved to: {citations_txt_path}")

    # Build citation keys
    cit_key_examples = build_citation_keys(citations)
    citation_keys = set(cit_key_examples.keys())
    log_lines.append(f"Number of distinct (author|year) citation keys: {len(citation_keys)}")

    # ------------------------------------------------------------------
    # 2) Extract reference list and parse entries
    # ------------------------------------------------------------------
    log_lines.append("\n[Step 2] Extracting reference list from DOCX...")

    doc = Document(docx_path)
    ref_lines = extract_reference_paragraphs(doc)
    log_lines.append(f"Number of reference lines found after reference heading: {len(ref_lines)}")

    if not ref_lines:
        log_lines.append("ERROR: No reference list detected (no suitable heading or nothing after it).")
        ref_entries = []
        ref_keys = set()
        missing_keys = set()
        unused_keys = set()
    else:
        ref_entries = parse_reference_entries(ref_lines)
        log_lines.append(f"Number of parsed reference entries (with year and first author): {len(ref_entries)}")

        ref_keys = {e["key"] for e in ref_entries}

        # ------------------------------------------------------------------
        # 3) Check consistency between in-text citations and reference list
        # ------------------------------------------------------------------
        log_lines.append("\n[Step 3] Checking citation–reference consistency...")

        missing_keys = citation_keys - ref_keys
        unused_keys = ref_keys - citation_keys

        # 3.1 In-text citation keys not found in reference list
        if missing_keys:
            log_lines.append("\n[ERROR] In-text citations with no matching reference entry (by first author + year):")
            for key in sorted(missing_keys):
                examples = "; ".join(sorted(cit_key_examples[key]))
                log_lines.append(f"  key = {key}   from citation(s): {examples}")
        else:
            log_lines.append("\n[OK] All detected in-text citation keys have at least one matching reference entry.")

        # 3.2 Reference entries that are never cited in the text
        if unused_keys:
            log_lines.append(
                "\n[INFO] Reference entries that were not detected in any in-text citation "
                "(not necessarily an error, but potential clean-up):"
            )
            for e in ref_entries:
                if e["key"] in unused_keys:
                    log_lines.append(f"  - {e['raw']}")
        else:
            log_lines.append("\n[OK] Every parsed reference entry is cited at least once in the text (by first author + year).")

    # ------------------------------------------------------------------
    # 4) Overall summary
    # ------------------------------------------------------------------
    log_lines.append("\n[Step 4] Overall summary")

    total_cit_parentheses = len(citations)
    total_cit_keys = len(citation_keys)
    total_ref_lines = len(ref_lines)
    total_ref_entries = len(ref_entries)
    total_missing = len(missing_keys)
    total_unused = len(unused_keys)

    log_lines.append(f"Total citation parentheses in text: {total_cit_parentheses}")
    log_lines.append(f"Total distinct citation keys (author|year): {total_cit_keys}")
    log_lines.append(f"Total reference lines after heading: {total_ref_lines}")
    log_lines.append(f"Total parsed reference entries: {total_ref_entries}")
    log_lines.append(f"Number of missing citation keys (no match in references): {total_missing}")
    log_lines.append(f"Number of unused reference entries (not cited in text): {total_unused}")

    if not ref_lines:
        log_lines.append(
            "\nOVERALL: Could not assess citation consistency because no reference list was detected."
        )
    else:
        if missing_keys:
            log_lines.append(
                "\nOVERALL: Citation–reference consistency has ERRORS. "
                "Some in-text citations (first author + year) do not have matching entries in the reference list. "
                "Please add or correct those references."
            )
        else:
            log_lines.append(
                "\nOVERALL: Citation–reference consistency looks GOOD. "
                "All detected in-text citation keys have at least one matching reference entry. "
                "Any unused references listed above are optional clean-up (not a formal style violation)."
            )

    # ------------------------------------------------------------------
    # 5) Write report to file
    # ------------------------------------------------------------------
    with report_txt_path.open("w", encoding="utf-8") as f:
        f.write("\n".join(log_lines))


if __name__ == "__main__":
    main()
